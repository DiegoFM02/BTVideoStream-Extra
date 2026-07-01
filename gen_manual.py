from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GUINDA = RGBColor(0x6B, 0x1A, 0x37)
AZUL   = RGBColor(0x14, 0x3C, 0x8C)
DARK   = RGBColor(0x11, 0x11, 0x11)
GRAY   = RGBColor(0x55, 0x55, 0x55)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
GREEN  = RGBColor(0x1B, 0x5E, 0x20)
RED    = RGBColor(0xB7, 0x1C, 0x1C)

doc = Document()
for s in doc.sections:
    s.top_margin    = Cm(2)
    s.bottom_margin = Cm(2)
    s.left_margin   = Cm(2.5)
    s.right_margin  = Cm(2.5)

# ── helpers ──────────────────────────────────────────────────────────────────

def sf(run, bold=False, size=11, color=DARK, italic=False, mono=False):
    run.bold = bold; run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = 'Courier New' if mono else 'Times New Roman'

def border_line(p, color='6B1A37', sz='8', pos='bottom'):
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    el   = OxmlElement(f'w:{pos}')
    el.set(qn('w:val'), 'single'); el.set(qn('w:sz'), sz)
    el.set(qn('w:space'), '1');    el.set(qn('w:color'), color)
    pBdr.append(el); pPr.append(pBdr)

def shading(p, fill='F4F4F4'):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill);   pPr.append(shd)

def cell_shade(cell, fill):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill);   tcPr.append(shd)

def h1(text, color=GUINDA):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text.upper())
    sf(r, bold=True, size=12, color=color)
    border_line(p, sz='10')
    return p

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    sf(r, bold=True, size=11, color=GUINDA)
    return p

def h3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    sf(r, bold=True, size=10.5, color=DARK)
    return p

def para(parts, space_after=4, justify=True):
    p = doc.add_paragraph()
    if justify: p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(space_after)
    if isinstance(parts, str):
        r = p.add_run(parts); sf(r, size=10.5)
    else:
        for item in parts:
            if isinstance(item, str):
                r = p.add_run(item); sf(r, size=10.5)
            else:
                t, bold = item[0], item[1]
                ital = item[2] if len(item) > 2 else False
                mn   = item[3] if len(item) > 3 else False
                col  = item[4] if len(item) > 4 else DARK
                r = p.add_run(t); sf(r, bold=bold, size=10.5, italic=ital, mono=mn, color=col)
    return p

def bullet(text, bold_pfx=None, indent=0.4, size=10.5, icon='•'):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(indent)
    if bold_pfx:
        r1 = p.add_run(f'{icon}  {bold_pfx}'); sf(r1, bold=True, size=size)
        r2 = p.add_run(text);                   sf(r2, size=size)
    else:
        r = p.add_run(f'{icon}  {text}');       sf(r, size=size)
    return p

def numbered(items, start=1, indent=0.4):
    for i, text in enumerate(items, start):
        p = doc.add_paragraph()
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.left_indent  = Cm(indent)
        if isinstance(text, str):
            r = p.add_run(f'{i}.  {text}'); sf(r, size=10.5)
        else:
            r0 = p.add_run(f'{i}.  '); sf(r0, size=10.5)
            for item in text:
                if isinstance(item, str):
                    r = p.add_run(item); sf(r, size=10.5)
                else:
                    t, bold = item[0], item[1]
                    mn = item[2] if len(item) > 2 else False
                    r  = p.add_run(t); sf(r, bold=bold, size=10.5, mono=mn)

def note(text, fill='FFF8E1', border='F9A825'):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.left_indent  = Cm(0.4)
    p.paragraph_format.right_indent = Cm(0.4)
    r = p.add_run(f'⚠  {text}'); sf(r, size=10, italic=True, color=RGBColor(0x4E, 0x34, 0x00))
    shading(p, fill='FFF3CD')

def tip(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(5)
    p.paragraph_format.left_indent  = Cm(0.4)
    p.paragraph_format.right_indent = Cm(0.4)
    r = p.add_run(f'✔  {text}'); sf(r, size=10, italic=True, color=GREEN)
    shading(p, fill='F1F8E9')

def codeblock(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.4)
    r = p.add_run(text); sf(r, size=8.5, mono=True)
    shading(p, fill='F4F4F4')

def add_table(headers, rows, col_widths=None, hdr_color='6B1A37'):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]; cell.text = ''
        r = cell.paragraphs[0].add_run(h); sf(r, bold=True, size=9, color=WHITE)
        cell_shade(cell, hdr_color)
    for ri, row_data in enumerate(rows):
        fill = 'FDF5F7' if ri % 2 == 1 else 'FFFFFF'
        for ci, ct in enumerate(row_data):
            cell = t.rows[ri+1].cells[ci]; cell.text = ''
            r = cell.paragraphs[0].add_run(ct); sf(r, size=9)
            cell_shade(cell, fill)
    if col_widths:
        for row in t.rows:
            for ci, w in enumerate(col_widths):
                row.cells[ci].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

def page_break():
    from docx.oxml import OxmlElement
    p = doc.add_paragraph()
    r = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r._r.append(br)

# ═══════════════════════════════════════════════════════════════════════════════
#  PORTADA
# ═══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
r = p.add_run('INSTITUTO POLITÉCNICO NACIONAL  ·  ESCUELA SUPERIOR DE CÓMPUTO')
sf(r, size=9, color=GRAY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run('BTVideoStream')
sf(r, bold=True, size=22, color=GUINDA)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
r = p.add_run('Manual Técnico y de Usuario')
sf(r, bold=True, size=14, color=DARK)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
r = p.add_run('Versión 1.0  ·  Julio 2026')
sf(r, size=10, color=GRAY, italic=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(10)
r = p.add_run('Desarrollo de Aplicaciones Móviles Nativas — Examen Extraordinario')
sf(r, size=10, color=GRAY)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(10)
border_line(p, sz='12')

# Descripción breve en portada
para([
    ('BTVideoStream', True),
    ' permite a un dispositivo Android ',
    ('sin conexión a Internet', True),
    ' buscar y reproducir videos de YouTube y TikTok a través de un segundo '
    'dispositivo ',
    ('con acceso a Internet', True),
    ', usando ',
    ('Bluetooth clásico (RFCOMM/SPP)', True),
    ' como canal de comunicación. El dispositivo con Internet actúa como '
    'Servidor; el dispositivo sin Internet actúa como Cliente.',
])

add_table(
    ['Característica', 'Valor'],
    [
        ['Plataforma',         'Android 7.0 (API 24) o superior'],
        ['Lenguaje',           'Kotlin — Android Nativo'],
        ['Interfaz',           'Jetpack Compose + Material 3'],
        ['Reproductor',        'AndroidX Media3 / ExoPlayer'],
        ['Comunicación',       'Bluetooth Clásico RFCOMM/SPP'],
        ['Fuentes de video',   'YouTube, TikTok, Google/DuckDuckGo'],
        ['Frameworks terceros','Ninguno'],
        ['Repositorio',        'github.com/DiegoFM02/BTVideoStream-Extra'],
    ],
    col_widths=[4.5, 9.1]
)

# ═══════════════════════════════════════════════════════════════════════════════
#  PARTE I — MANUAL DE USUARIO
# ═══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(12)
p.paragraph_format.space_after  = Pt(6)
border_line(p, sz='16')
r = p.add_run('PARTE I — MANUAL DE USUARIO')
r2 = p.runs[0]  # mismo run
p2 = doc.paragraphs[-1]
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sf(r, bold=True, size=13, color=GUINDA)

# ── 1. REQUISITOS ──────────────────────────────────────────────────────────────
h1('1. Requisitos del Sistema')

h2('1.1 Requisitos de Hardware')
add_table(
    ['Componente', 'Servidor (Dispositivo A)', 'Cliente (Dispositivo B)'],
    [
        ['Sistema operativo', 'Android 7.0+ (API 24)',         'Android 7.0+ (API 24)'],
        ['Bluetooth',         'Clásico 4.1 o superior',         'Clásico 4.1 o superior'],
        ['Conexión a Internet','Wi-Fi o datos móviles requerido','Desactivada (Wi-Fi y datos)'],
        ['Almacenamiento libre','≥ 500 MB (caché de videos)',    '≥ 200 MB (videos temporales)'],
        ['RAM recomendada',   '≥ 2 GB',                         '≥ 2 GB'],
    ],
    col_widths=[3.8, 4.9, 4.9]
)
note('La app requiere dos dispositivos físicos Android. Los emuladores no soportan Bluetooth.')

h2('1.2 Permisos Requeridos')
add_table(
    ['Permiso', 'Propósito', 'Obligatorio'],
    [
        ['BLUETOOTH_CONNECT (API 31+)',   'Conectar y desconectar dispositivos BT', 'Sí'],
        ['BLUETOOTH_SCAN (API 31+)',       'Descubrir dispositivos cercanos',         'Sí'],
        ['BLUETOOTH_ADVERTISE (API 31+)', 'Anunciar el servidor BT',                 'Sí'],
        ['BLUETOOTH / BLUETOOTH_ADMIN',  'Bluetooth en Android ≤ 11',               'Sí (API ≤ 30)'],
        ['ACCESS_FINE_LOCATION',          'Requerido para BT discovery en API ≤ 30', 'Sí (API ≤ 30)'],
        ['POST_NOTIFICATIONS (API 33+)',  'Notificaciones de estado de conexión',    'Recomendado'],
        ['INTERNET',                       'Solo el Servidor accede a Internet',      'Sí (Servidor)'],
    ],
    col_widths=[4.5, 6.5, 2.6]
)

# ── 2. INSTALACIÓN ─────────────────────────────────────────────────────────────
h1('2. Instalación')

h2('2.1 Obtener el APK')
para('Descarga el archivo APK desde el repositorio del proyecto o compílalo desde el código fuente:')
codeblock(
    'git clone https://github.com/DiegoFM02/BTVideoStream-Extra.git\n'
    'cd BTVideoStream-Extra\n'
    './gradlew assembleDebug\n'
    '# APK generado en: app/build/outputs/apk/debug/app-debug.apk'
)

h2('2.2 Instalar en los dispositivos')
note('Activa "Fuentes desconocidas" en Ajustes → Seguridad antes de instalar el APK manualmente.')
numbered([
    'Transfiere el archivo app-debug.apk a cada dispositivo (cable USB, AirDrop, Drive, etc.).',
    'En el dispositivo, abre el archivo APK desde el explorador de archivos.',
    'Acepta la instalación cuando el sistema lo solicite.',
    'Repite el proceso en el segundo dispositivo.',
    'Empareja ambos dispositivos por Bluetooth desde los Ajustes del sistema antes de abrir la app.',
])

h2('2.3 Emparejamiento Bluetooth previo')
para('Antes de usar la aplicación, los dos dispositivos deben estar emparejados entre sí:')
numbered([
    'En ambos dispositivos: Ajustes → Bluetooth → Activar.',
    'En uno de ellos: tocar "Buscar dispositivos" o "Vincular nuevo dispositivo".',
    'Seleccionar el otro dispositivo en la lista y confirmar el PIN en ambos.',
    'Una vez emparejados, no es necesario repetir este paso en sesiones futuras.',
])

# ── 3. PRIMER USO ──────────────────────────────────────────────────────────────
h1('3. Primer Uso — Permisos')
para('Al abrir la aplicación por primera vez, el sistema solicitará los permisos de Bluetooth. '
     'Es necesario conceder todos los permisos para que la app funcione correctamente.')
numbered([
    'Abrir la app BTVideoStream en el dispositivo.',
    'Aparecerá una pantalla de permisos. Tocar "Conceder permisos".',
    'Aceptar cada permiso solicitado por el sistema Android.',
    ['En Android 13+: aceptar también el permiso de ', ('Notificaciones', True), ' cuando se solicite.'],
    'Una vez concedidos todos los permisos, la app avanzará a la pantalla de selección de rol.',
])
tip('Si accidentalmente denegas un permiso, ve a Ajustes → Apps → BTVideoStream → Permisos para activarlos manualmente.')

# ── 4. SELECCIÓN DE ROL ────────────────────────────────────────────────────────
h1('4. Selección de Rol')
para('Cada vez que se abre la app, se presenta la pantalla de selección de rol. '
     'Ambos dispositivos deben abrir la app y elegir roles distintos.')

add_table(
    ['Rol', 'Cuándo elegirlo', 'Requisito'],
    [
        ['🖥  Soy Servidor',
         'Dispositivo con Wi-Fi o datos móviles activos',
         'Conexión a Internet activa'],
        ['📱  Soy Cliente',
         'Dispositivo sin Internet (Wi-Fi y datos desactivados)',
         'Solo Bluetooth habilitado'],
    ],
    col_widths=[3.2, 6.0, 4.4]
)
note('Primero debe iniciarse el Servidor y luego conectarse el Cliente, nunca al revés.')

# ── 5. USO COMO SERVIDOR ───────────────────────────────────────────────────────
h1('5. Uso como Servidor')

h2('5.1 Iniciar el servidor')
numbered([
    'Seleccionar "🖥  Soy Servidor" en la pantalla de roles.',
    'Tocar el botón "Esperar conexión".',
    'El servidor comenzará a escuchar conexiones Bluetooth entrantes.',
    'El indicador de estado cambiará a color naranja pulsante ("Esperando cliente…").',
    'Esperar a que el Cliente se conecte.',
])
tip('Una vez conectado el cliente, el indicador cambia a verde. El log de actividad muestra en tiempo real cada búsqueda y descarga procesada.')

h2('5.2 Pantalla del Servidor')
add_table(
    ['Elemento de UI', 'Función'],
    [
        ['Indicador de estado (barra superior)', 'Muestra si hay cliente conectado (verde), esperando (naranja) o sin conexión (rojo)'],
        ['Nombre del dispositivo conectado',     'Aparece junto al indicador cuando hay un cliente activo'],
        ['Log de actividad',                     'Lista cronológica de todas las acciones: búsquedas, descargas, errores y transferencias'],
        ['Botón de tema (esquina superior)',      'Cambia entre tema Guinda (IPN) y Azul (ESCOM), y alterna modo claro/oscuro'],
    ],
    col_widths=[5.0, 8.6]
)
note('El servidor NO requiere interacción manual mientras opera. Todas las solicitudes son procesadas automáticamente al recibirlas del cliente.')

# ── 6. USO COMO CLIENTE ────────────────────────────────────────────────────────
h1('6. Uso como Cliente')

h2('6.1 Conectarse al Servidor')
numbered([
    'Seleccionar "📱  Soy Cliente" en la pantalla de roles.',
    'La app mostrará la lista de dispositivos Bluetooth emparejados.',
    'Tocar el nombre del dispositivo que actúa como Servidor.',
    'El indicador cambiará a naranja ("Conectando…") y luego a verde ("Conectado").',
    'Una notificación push confirmará la conexión exitosa.',
])
note('Si la conexión falla, la app reintentará automáticamente hasta 5 veces. Si persiste, verifica que el Servidor esté en modo espera.')

h2('6.2 Buscar un video')
numbered([
    ['Seleccionar la fuente de búsqueda:', ('YouTube', True), ', ', ('TikTok', True), ' o ', ('Google', True), '.'],
    'Escribir el título, palabras clave o nombre del artista en la barra de búsqueda.',
    'Tocar el ícono de búsqueda o presionar Enter.',
    'Esperar los resultados (el servidor los procesa y envía por Bluetooth).',
    'La lista de resultados mostrará: miniatura (URL), título, canal y duración.',
])

h2('6.3 Reproducir un video')
numbered([
    'Tocar cualquier resultado de la lista para solicitarlo al servidor.',
    'Seleccionar la calidad deseada en el diálogo que aparece.',
    'El servidor descargará y transferirá el video por Bluetooth (puede tardar algunos segundos).',
    'El video se reproducirá automáticamente con el reproductor integrado.',
])

add_table(
    ['Calidad', 'Resolución', 'Cuándo usarla'],
    [
        ['Baja calidad',   '360p',       'Conexión BT lenta o video largo (recomendado para pruebas)'],
        ['Alta calidad',   '720p',       'Videos cortos o cuando se requiere mayor definición'],
        ['Solo audio',     'Audio only', 'Modo de bajo consumo — menor tiempo de transferencia'],
    ],
    col_widths=[3.0, 2.6, 8.0]
)

h2('6.4 Controles del reproductor')
add_table(
    ['Control', 'Función'],
    [
        ['▶ / ⏸  Reproducir/Pausar', 'Inicia o pausa la reproducción del video'],
        ['⏪ / ⏩  Retroceder/Adelantar', 'Salta 10 segundos hacia atrás o adelante'],
        ['Barra de progreso (seek)', 'Arrastra para navegar a cualquier punto del video'],
        ['Pantalla completa', 'Rota el dispositivo para ver en horizontal'],
    ],
    col_widths=[5.2, 8.4]
)

h2('6.5 Historial y funciones adicionales')
bullet('El historial de videos reproducidos se guarda automáticamente y persiste entre sesiones.')
bullet('El modo de privacidad permite reproducir un video sin registrarlo en el historial.')
bullet('El servidor aplica caché: si el mismo video se solicita de nuevo, la transferencia es casi instantánea.')

# ── 7. TEMAS Y MODO OSCURO ─────────────────────────────────────────────────────
h1('7. Personalización — Temas y Modo Oscuro')
para('La aplicación soporta dos temas y tres modos de iluminación, configurables desde el ícono '
     'de paleta en la esquina superior derecha de cualquier pantalla.')

add_table(
    ['Opción', 'Descripción'],
    [
        ['Tema Guinda (IPN)',   'Esquema de colores borgoña/guinda, color representativo del IPN'],
        ['Tema Azul (ESCOM)',   'Esquema de colores azul, color representativo de la ESCOM'],
        ['Modo claro',          'Fondo blanco, texto oscuro'],
        ['Modo oscuro',         'Fondo oscuro, texto claro — reduce fatiga visual'],
        ['Modo automático',     'Sigue la configuración del sistema operativo del dispositivo'],
    ],
    col_widths=[3.8, 9.8]
)
tip('La preferencia de tema se guarda automáticamente. No es necesario reconfigurar al reiniciar la app.')

# ── 8. RESOLUCIÓN DE PROBLEMAS ─────────────────────────────────────────────────
h1('8. Resolución de Problemas')
add_table(
    ['Problema', 'Causa probable', 'Solución'],
    [
        ['"No se encontraron dispositivos"',
         'El dispositivo Servidor no está emparejado',
         'Emparejar ambos dispositivos desde Ajustes → Bluetooth del sistema'],
        ['La conexión falla repetidamente',
         'El Servidor aún no está en modo espera',
         'Asegurarse de tocar "Esperar conexión" en el Servidor antes de conectar el Cliente'],
        ['"Servidores ocupados" en YouTube',
         'El CDN de YouTube rechaza la descarga (IP binding)',
         'Intentar con otro video, o usar TikTok que funciona de forma estable'],
        ['TikTok no devuelve resultados',
         'Sin conexión a Internet en el Servidor',
         'Verificar que el Servidor tenga Wi-Fi o datos activos'],
        ['Video se ve cortado o con errores',
         'Interferencia Bluetooth durante la transferencia',
         'Acercar los dispositivos y volver a solicitar el video (se usa la caché)'],
        ['La app no muestra los permisos',
         'Permisos previamente denegados',
         'Ajustes → Apps → BTVideoStream → Permisos → Activar todos los de Bluetooth'],
        ['Transferencia muy lenta',
         'Interferencia en la banda de 2.4 GHz',
         'Alejar el dispositivo de routers Wi-Fi y otros dispositivos BT simultáneos'],
    ],
    col_widths=[3.8, 4.2, 5.6]
)

# ═══════════════════════════════════════════════════════════════════════════════
#  PARTE II — MANUAL TÉCNICO
# ═══════════════════════════════════════════════════════════════════════════════

page_break()

p = doc.add_paragraph()
p.paragraph_format.space_after  = Pt(6)
border_line(p, sz='16')
r = p.add_run('PARTE II — MANUAL TÉCNICO')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sf(r, bold=True, size=13, color=GUINDA)

# ── 9. ESTRUCTURA DEL PROYECTO ─────────────────────────────────────────────────
h1('9. Estructura del Proyecto')
codeblock(
    'app/src/main/java/com/example/btvideostream/\n'
    '│\n'
    '├── MainActivity.kt                # Entry point, navegación y permisos\n'
    '│\n'
    '├── bluetooth/\n'
    '│   ├── BluetoothManager.kt        # RFCOMM server/client, reconexión, chunking\n'
    '│   ├── ConnectionState.kt         # Sealed class: Disconnected/Connecting/Connected/Error\n'
    '│   └── protocol/\n'
    '│       ├── Message.kt             # Tipos de mensaje (SearchRequest, VideoChunk, etc.)\n'
    '│       └── MessageSerializer.kt   # Serialización/deserialización binaria\n'
    '│\n'
    '├── core/permissions/\n'
    '│   └── BluetoothPermissions.kt    # Permisos BT por versión de Android\n'
    '│\n'
    '├── notifications/\n'
    '│   └── BtNotificationManager.kt   # Notificaciones push de estado BT\n'
    '│\n'
    '├── server/\n'
    '│   ├── Models.kt                  # VideoResult, StreamInfo, WebResult\n'
    '│   ├── YouTubeService.kt          # Búsqueda y extracción de stream de YouTube\n'
    '│   ├── TikTokService.kt           # Búsqueda y extracción de video de TikTok\n'
    '│   ├── GoogleService.kt           # Búsqueda web via DuckDuckGo\n'
    '│   └── VideoRepository.kt        # Caché, descarga y envío BT\n'
    '│\n'
    '└── ui/\n'
    '    ├── screens/                   # RoleSelectionScreen, ServerScreen, ClientScreen\n'
    '    ├── client/ClientViewModel.kt  # Lógica del cliente\n'
    '    ├── server/ServerViewModel.kt  # Lógica del servidor\n'
    '    ├── components/                # ConnectionIndicator, ThemeSwitcher\n'
    '    └── theme/                     # Color.kt, Theme.kt, ThemeController.kt, Type.kt'
)

# ── 10. MÓDULO BLUETOOTH ───────────────────────────────────────────────────────
h1('10. Módulo Bluetooth — BluetoothManager')

h2('10.1 Protocolo de comunicación')
para([
    'El protocolo usa un framing binario sobre el socket RFCOMM. Cada mensaje tiene una cabecera '
    'de ',
    ('5 bytes', True),
    ' seguida del payload variable:',
])
codeblock(
    '┌──────────┬──────────────────────────┬──────────────────────────────────────┐\n'
    '│  type    │        length            │           payload                    │\n'
    '│  1 byte  │   4 bytes (Big-Endian)   │         N bytes (máx. 20 MB)         │\n'
    '└──────────┴──────────────────────────┴──────────────────────────────────────┘'
)

h2('10.2 Tipos de mensaje y valores')
add_table(
    ['Tipo', 'Hex', 'Payload', 'Dir.'],
    [
        ['PING / PONG',    '0x01/0x02', 'Vacío',                                     'Ambas'],
        ['DISCONNECT',     '0x03',       'Vacío',                                     'Ambas'],
        ['SEARCH_REQUEST', '0x10',       'UTF-8: "fuente|consulta"',                  'C→S'],
        ['SEARCH_RESULTS', '0x11',       'UTF-8: JSON array de resultados',           'S→C'],
        ['VIDEO_REQUEST',  '0x20',       '[quality: 1 byte][videoId: UTF-8]',         'C→S'],
        ['VIDEO_CHUNK',    '0x21',       '[index: 4 bytes BE] + bytes del fragmento', 'S→C'],
        ['VIDEO_END',      '0x22',       'Vacío — señal de fin de transferencia',     'S→C'],
        ['VIDEO_ERROR',    '0x23',       'UTF-8: mensaje de error',                   'S→C'],
        ['STATUS',         '0x30',       'UTF-8: mensaje / keepalive',                'Ambas'],
    ],
    col_widths=[3.0, 1.5, 7.1, 2.0]
)

h2('10.3 Gestión de conexiones')
bullet('Llama a listenUsingRfcommWithServiceRecord() con un UUID de servicio fijo '
       'y queda en espera de conexiones entrantes en un hilo IO.', bold_pfx='Servidor: ')
bullet('Llama a createRfcommSocketToServiceRecord() con el mismo UUID y ejecuta '
       'hasta 5 reintentos con espera progresiva (500 ms × intento) ante fallos.', bold_pfx='Cliente: ')
bullet('Todas las escrituras al socket pasan por un Channel<Message> FIFO. '
       'Un único consumidor drena la cola en orden, evitando condiciones de carrera.', bold_pfx='Escritura concurrente: ')
bullet('El servidor envía un mensaje STATUS cada 2 segundos para mantener el enlace '
       'activo y detectar desconexiones silenciosas.', bold_pfx='Keepalive: ')
bullet('Los videos se dividen en fragmentos de 2 048 bytes. Cada chunk lleva '
       'un índice de 4 bytes para reordenamiento en el cliente.', bold_pfx='Chunking: ')

# ── 11. MÓDULO SERVIDOR ────────────────────────────────────────────────────────
h1('11. Módulos del Servidor')

h2('11.1 YouTubeService')
add_table(
    ['Método', 'Descripción'],
    [
        ['search(query, maxResults)',
         'Obtiene youtube.com/results con User-Agent de navegador y parsea ytInitialData '
         'del HTML para extraer título, canal y miniatura de hasta 15 resultados.'],
        ['getStreamUrl(videoId, preferLow)',
         'Intenta 3 estrategias en orden: (1) parseo de ytInitialPlayerResponse del HTML '
         'de la página watch, (2) Cobalt API, (3) Invidious con local=true.'],
        ['extractFromWatchPage()',
         'Parsea el HTML de youtube.com/watch para extraer URLs directas (muxed) '
         'del array formats. Captura cookies de sesión para el CDN.'],
    ],
    col_widths=[4.0, 9.6]
)
note('Se fuerza IPv4 (java.net.preferIPv4Stack=true) para que el parámetro ip= firmado en la URL del CDN coincida con la dirección de descarga.')

h2('11.2 TikTokService')
add_table(
    ['Método', 'Descripción'],
    [
        ['search(query, maxResults)',
         'POST a tikwm.com/api/feed/search con Content-Type: application/x-www-form-urlencoded. '
         'Los IDs de video se prefijan con "tt_" para distinguirlos de YouTube.'],
        ['getVideoUrl(videoId)',
         'GET a tikwm.com/api/?url=... para obtener la URL directa del video MP4.'],
    ],
    col_widths=[4.0, 9.6]
)
para([
    'TikWM usa una cadena SSL incompleta; se aplica un ',
    ('X509TrustManager', True, False, True),
    ' permisivo y un ',
    ('HostnameVerifier', True, False, True),
    ' personalizado exclusivamente para estas conexiones.',
])

h2('11.3 VideoRepository — Caché y Descarga')
bullet('Cada video descargado se almacena en cacheDir/video_cache/{videoId}.mp4. '
       'En solicitudes posteriores se sirve desde disco, sin volver a descargar.', bold_pfx='Caché: ')
bullet('downloadUrl() descarga los bytes completos del video con User-Agent de '
       'navegador, cabeceras Referer/Origin y cookies de sesión para YouTube.', bold_pfx='Descarga: ')
bullet('Los bytes descargados se dividen en chunks de 2 048 bytes. Cada chunk '
       'se envía como VIDEO_CHUNK con su índice. Al terminar, se envía VIDEO_END.', bold_pfx='Envío: ')

# ── 12. MÓDULO CLIENTE ─────────────────────────────────────────────────────────
h1('12. Módulo Cliente')

h2('12.1 ClientViewModel')
add_table(
    ['Componente', 'Descripción'],
    [
        ['chunkBuffer: Map<Int, ByteArray>',
         'Almacena chunks recibidos indexados. Permite reordenarlos si llegan fuera de orden.'],
        ['Recepción de VIDEO_CHUNK',
         'Cada chunk recibido se agrega al buffer por su índice.'],
        ['Recepción de VIDEO_END',
         'Concatena todos los chunks en orden, escribe received_video.mp4 y lanza ExoPlayer.'],
        ['Recepción de SEARCH_RESULTS',
         'Parsea el JSON y actualiza el StateFlow<List<VideoResult>> que observa la UI.'],
        ['requestOpenUrl()',
         'Envía STATUS("OPEN_URL:https://...") al servidor para abrir URLs externas.'],
    ],
    col_widths=[4.5, 9.1]
)

h2('12.2 Reproducción con ExoPlayer')
bullet('El archivo MP4 recibido se pasa a ExoPlayer como FileMediaSource desde cacheDir.')
bullet('Soporta reproducir/pausar, seek, retroceder 10 s, adelantar 10 s y pantalla completa.')
bullet('ExoPlayer gestiona internamente el buffer y la decodificación de hardware.')

# ── 13. COMPATIBILIDAD ─────────────────────────────────────────────────────────
h1('13. Compatibilidad')
add_table(
    ['Versión Android', 'API', 'Soporte', 'Notas'],
    [
        ['Android 7.0 Nougat',    '24', '✔ Soportado', 'Permisos BT legacy (BLUETOOTH + ADMIN)'],
        ['Android 8.0–9.0',       '26–28', '✔ Soportado', 'Sin cambios relevantes'],
        ['Android 10–11',         '29–30', '✔ Soportado', 'ACCESS_FINE_LOCATION requerido para BT scan'],
        ['Android 12–12L',        '31–32', '✔ Soportado', 'Nuevos permisos BLUETOOTH_CONNECT/SCAN/ADVERTISE'],
        ['Android 13',            '33', '✔ Soportado', 'POST_NOTIFICATIONS requerido para notificaciones'],
        ['Android 14–15',         '34–35', '✔ Soportado', 'Probado como target SDK'],
        ['Android 16 (API 36)',   '36',    '✔ Soportado', 'compileSdk y targetSdk actuales'],
    ],
    col_widths=[3.8, 1.5, 2.8, 5.5]
)

h2('Formatos de video soportados')
add_table(
    ['Fuente', 'Formato', 'Resoluciones', 'Limitación'],
    [
        ['YouTube',  'MP4 (muxed)', '360p, 720p',   'Solo formatos combinados. DASH no soportado.'],
        ['TikTok',   'MP4',         'Original',      'Dependiente de la API de TikWM.'],
        ['Google',   'N/A (links)', 'N/A',           'Redirige al navegador; no descarga video.'],
    ],
    col_widths=[2.4, 2.4, 2.8, 6.0]
)

# ── 14. LIMITACIONES ───────────────────────────────────────────────────────────
h1('14. Restricciones y Limitaciones Conocidas')
add_table(
    ['Limitación', 'Detalle'],
    [
        ['YouTube HTTP 403',
         'El CDN de YouTube firma la URL con la IP del cliente y, desde 2024, requiere un '
         'poToken generado por JavaScript. Sin motor JS nativo, la descarga puede ser '
         'rechazada. TikTok funciona de forma estable como alternativa.'],
        ['Solo formatos muxed en YouTube',
         'Los streams DASH (adaptivos) separan video y audio. Combinarlos requiere FFmpeg, '
         'no disponible sin librerías de terceros.'],
        ['Un cliente simultáneo',
         'El servidor acepta una única conexión RFCOMM activa a la vez.'],
        ['Velocidad de transferencia BT',
         'RFCOMM ofrece ~2–3 Mbps teóricos. Videos de 360p (~30–60 MB) tardan '
         '20–60 segundos. La caché elimina este tiempo en solicitudes repetidas.'],
        ['Alcance Bluetooth',
         'Distancia efectiva de ~10 m en interiores sin obstáculos. La tasa de error '
         'aumenta con la distancia o la presencia de paredes gruesas.'],
        ['Internet en el Servidor',
         'Si el Servidor pierde la conexión a Internet durante una descarga, '
         'el video fallará. La caché protege las descargas ya completadas.'],
    ],
    col_widths=[4.0, 9.6]
)

# ── 15. GLOSARIO ───────────────────────────────────────────────────────────────
h1('15. Glosario')
add_table(
    ['Término', 'Definición'],
    [
        ['RFCOMM',          'Radio Frequency Communication — protocolo de emulación de puerto serie sobre Bluetooth'],
        ['SPP',             'Serial Port Profile — perfil Bluetooth que usa RFCOMM para comunicación bidireccional'],
        ['Muxed',           'Formato de video con audio y video combinados en un único archivo/stream'],
        ['DASH',            'Dynamic Adaptive Streaming over HTTP — streams adaptivos con audio y video separados'],
        ['ytInitialPlayerResponse', 'Objeto JSON incrustado en el HTML de YouTube con metadatos y URLs del video'],
        ['poToken',         'Proof-of-Origin Token — token criptográfico de YouTube que requiere ejecución de JS'],
        ['CDN',             'Content Delivery Network — red de servidores distribuidos que sirve el video de YouTube'],
        ['ExoPlayer',       'Reproductor de medios de Android (AndroidX Media3) con soporte para múltiples formatos'],
        ['Chunk',           'Fragmento de datos de tamaño fijo (2 048 bytes) para transferencia por Bluetooth'],
        ['Caché',           'Almacenamiento local de videos descargados para evitar re-descargas'],
        ['MVVM',            'Model-View-ViewModel — patrón de arquitectura que separa lógica de negocio de la UI'],
        ['StateFlow',       'Flujo de Kotlin Coroutines que emite el último estado y notifica a los observadores'],
    ],
    col_widths=[4.2, 9.4]
)

# ── PIE ────────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(10)
border_line(p, color='CCCCCC', sz='4', pos='top')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    'BTVideoStream v1.0  ·  ESCOM – IPN  ·  Desarrollo de Aplicaciones Móviles Nativas  ·  Julio 2026\n'
    'github.com/DiegoFM02/BTVideoStream-Extra'
)
sf(r, size=8, color=GRAY, italic=True)

doc.save(r'D:\Moviles\Tareas\BTVideoStream\manual_tecnico_usuario.docx')
print('Manual generado correctamente.')
